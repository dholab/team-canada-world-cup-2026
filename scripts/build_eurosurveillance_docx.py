#!/usr/bin/env python3
"""Build the Word files Eurosurveillance's submission system requires.

Eurosurveillance reviews anonymised manuscripts, so the submission is split
across separate files rather than the single PDF used for the preprint:

  Title_page.docx        authors, affiliations, correspondence, word counts
  Manuscript_anon.docx   the manuscript with line numbers and NO author-
                         identifiable information
  Key_public_health_message.docx   their required PH-message template content

"pdf files cannot be evaluated", so these are .docx. The preprint build is
untouched: it keeps the full author list and every statement.

What is removed from the anonymised text, per their instructions ("authors'
names, affiliations and contributions, as well as any acknowledgements"):

  - Authors' contributions, Acknowledgements, Conflict of interest and Funding
    statements. The last two are declared in the submission system instead, and
    naming the funders would identify the group.
  - Institution names in Methods (the IRB and the sequencing core), replaced
    with neutral wording.
  - The repository, LabKey and preprint URLs in Data availability, which resolve
    to the authors' own org and host. The statement is kept but neutralised, so
    the reviewer still learns the data are public.

Run:  python scripts/build_eurosurveillance_docx.py
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

import docx
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).resolve().parent.parent
# Built into a staging directory, not straight into the package folder:
# scripts/build_submission_package.py clears the package folder before
# copying into it, which would delete these files before use.
OUT = ROOT / "submission" / "_docx"

# Sections dropped wholesale from the anonymised manuscript.
DROP_SECTIONS = {
    "Authors' contributions",
    "Acknowledgements",
    "Conflict of interest",
    "Funding statement",
}

# Identifying strings in the retained text -> neutral replacements.
DEIDENTIFY = [
    (r"The University of Wisconsin[–-]Madison Institutional Review Board determined",
     "The institutional review board of the coordinating institution determined"),
    (r"The University of Wisconsin-Madison Institutional Review Board has determined",
     "The institutional review board of the coordinating institution has determined"),
    (r"This sequencing was done by the UW-Madison Biotechnology Center\.",
     "Sequencing was performed by a university core sequencing facility."),
]

# The Data availability statement is replaced entirely: every URL in it points
# at the authors' own GitHub organisation and institutional host.
DATA_AVAILABILITY_ANON = (
    "The data and analysis code for this study, including the per-cartridge "
    "results table behind every figure, are openly available in a public "
    "repository; the URL is withheld here because it identifies the authors and "
    "is given on the title page. Sequencing reads are deposited in the NCBI "
    "Sequence Read Archive under a BioProject accession, also given on the title "
    "page. Both datasets include two additional air samples not discussed in the "
    "manuscript: during part of the tournament two further samplers were tested "
    "for norovirus on the GeneXpert, and the two that were running at the time of "
    "the team's elimination are included in the sequencing datasets. No air "
    "sample tested positive for norovirus, and because coverage of the tournament "
    "was incomplete these were not included in the manuscript's primary dataset."
)

INLINE_MD = re.compile(r"\[\[?(\d+)\]\([^)]*\)\]?")      # citation links
LINK_MD = re.compile(r"\[([^\]]*)\]\([^)]*\)")             # other links
BOLD_SPLIT = re.compile(r"(\*\*.+?\*\*)", re.S)


def clean(text: str) -> str:
    """Markdown -> plain text, keeping citation numbers as [n]."""
    text = INLINE_MD.sub(lambda m: f"[{m.group(1)}]", text)
    text = LINK_MD.sub(r"\1", text)
    text = re.sub(r"\{\{<.*?>\}\}", "", text)
    text = text.replace("\\", "")
    return re.sub(r"[ \t]+", " ", text).strip()


def add_line_numbers(section) -> None:
    """Continuous line numbers, which Eurosurveillance requires."""
    ln = OxmlElement("w:lnNumType")
    ln.set(qn("w:countBy"), "1")
    ln.set(qn("w:start"), "1")
    ln.set(qn("w:restart"), "continuous")
    section._sectPr.append(ln)


def new_doc(line_numbers: bool = False) -> docx.Document:
    doc = docx.Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    if line_numbers:
        add_line_numbers(doc.sections[0])
    # Word stores an author name in the file's core properties. python-docx
    # defaults it to "python-docx", but set it explicitly so an anonymised
    # manuscript can never carry a real name in its metadata.
    cp = doc.core_properties
    cp.author = "Anonymous"
    cp.last_modified_by = "Anonymous"
    cp.comments = ""
    return doc


def write_para(doc, text: str, bold_lead: bool = True) -> None:
    """Add a paragraph, honouring **bold** runs."""
    p = doc.add_paragraph()
    for chunk in BOLD_SPLIT.split(text):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            p.add_run(chunk[2:-2]).bold = True
        else:
            p.add_run(chunk)


def parse_sections(md: str) -> list[tuple[int, str, list[str]]]:
    """[(level, heading, [paragraphs]), ...] from the synced prose."""
    md = re.sub(r"<!--.*?-->", "", md, flags=re.S)
    out: list[tuple[int, str, list[str]]] = []
    cur: tuple[int, str] | None = None
    buf: list[str] = []
    for block in re.split(r"\n\s*\n", md):
        block = block.strip()
        if not block:
            continue
        m = re.match(r"^(#{2,3})\s+(.*)$", block)
        if m:
            if cur:
                out.append((cur[0], cur[1], buf))
            cur, buf = (len(m.group(1)), m.group(2).strip()), []
        elif cur:
            buf.append(block)
    if cur:
        out.append((cur[0], cur[1], buf))
    return out


def word_count(sections) -> int:
    n = 0
    for lvl, head, paras in sections:
        if head in ("Introduction", "Methods", "Results", "Discussion"):
            pass
        n += 0
    return n


def build_manuscript(sections, title: str) -> Path:
    doc = new_doc(line_numbers=True)
    doc.add_paragraph(title).runs[0].bold = True
    doc.add_paragraph()

    skip = False
    for lvl, head, paras in sections:
        if head in DROP_SECTIONS:
            skip = True
            continue
        skip = False
        # The end-statements wrapper heading is BJSM-era scaffolding.
        if head == "Required end statements":
            continue
        if head == "Key public health message":
            continue  # supplied as its own file
        h = doc.add_paragraph()
        r = h.add_run(head)
        r.bold = True
        if lvl == 3:
            r.italic = True
        for para in paras:
            text = clean(para)
            if head == "Data availability":
                text = DATA_AVAILABILITY_ANON
            else:
                for pat, rep in DEIDENTIFY:
                    text = re.sub(pat, rep, text)
            if not text:
                continue
            write_para(doc, text)
            if head == "Data availability":
                break

    # References. Eurosurveillance requires a minimum of 15 for regular
    # articles, and an anonymised manuscript with no bibliography would be
    # returned. The entries are the authors' own published work in places, but
    # that is unavoidable and expected in any anonymised submission.
    h = doc.add_paragraph()
    h.add_run("References").bold = True
    refs = (ROOT / "_references.md").read_text()
    refs = re.sub(r"<!--.*?-->", "", refs, flags=re.S)
    for line in refs.splitlines():
        line = line.strip()
        if not re.match(r"^\d+\. ", line):
            continue
        para = doc.add_paragraph(clean(line))
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

    dest = OUT / "Manuscript_anon.docx"
    doc.save(dest)
    return dest


def build_ph_message(sections) -> Path:
    doc = new_doc()
    doc.add_paragraph("Key public health message").runs[0].bold = True
    doc.add_paragraph()
    for lvl, head, paras in sections:
        if head == "Key public health message":
            for para in paras:
                write_para(doc, clean(para))
            break
    dest = OUT / "Key_public_health_message.docx"
    doc.save(dest)
    return dest


def build_title_page(title: str, frontmatter: str, counts: dict) -> Path:
    doc = new_doc()

    def head(t):
        p = doc.add_paragraph()
        p.add_run(t).bold = True

    head("Title")
    doc.add_paragraph(title)

    head("Running head")
    doc.add_paragraph("Air sampling in a national team during the World Cup")

    fm = re.sub(r"<!--.*?-->", "", frontmatter, flags=re.S)
    for label in ("Authors", "Affiliations", "ORCID iDs", "Corresponding author"):
        m = re.search(rf"\*\*{label}\*\*\s*\n+(.*?)(?=\n\*\*|\Z)", fm, re.S)
        if not m:
            continue
        head(label)
        for line in m.group(1).strip().split("\n"):
            # "<sup>&amp;2</sup>" is the equal-contribution marker; render it
            # as a plain "&,2" rather than the raw "&2", which reads as a typo.
            line = clean(re.sub(r"<sup>(.*?)</sup>", r"\1", line)
                           .replace("&amp;", "&"))
            line = re.sub(r"^[-*]\s*", "", line)
            if line:
                doc.add_paragraph(line)

    head("Keywords")
    doc.add_paragraph("; ".join(counts["keywords"]))

    head("Word counts")
    doc.add_paragraph(f"Abstract: {counts['abstract']} words")
    doc.add_paragraph(f"Main text (Introduction to Conclusion): {counts['main']} words")
    doc.add_paragraph(f"References: {counts['refs']}")
    doc.add_paragraph("Figures: 4; Supplementary figures: 1; Tables: 0")

    head("Funding")
    doc.add_paragraph("This work was supported by Inkfish LLC and Heart of Racing.")

    head("Conflict of interest")
    doc.add_paragraph(
        "D.H.O. and S.L.O. are managing partners of Pathogenuity LLC, a "
        "consultancy that advises on topics including environmental monitoring "
        "for pathogens. D.H.O. and S.L.O. are Honorary professorial fellows at "
        "the University of Melbourne, Australia. The remaining authors have "
        "declared no conflicts of interest.")

    head("Authors' contributions")
    doc.add_paragraph(counts["contributions"])

    head("Acknowledgements")
    doc.add_paragraph(counts["acknowledgements"])

    head("Data availability")
    doc.add_paragraph(counts["data_availability"])

    head("Preprint")
    doc.add_paragraph(
        "A preprint of this manuscript is deposited on medRxiv, "
        "DOI 10.64898/2026.08.16.26360542. An interactive version is available "
        "at https://dholab.github.io/team-canada-world-cup-2026/.")

    dest = OUT / "Title_page.docx"
    doc.save(dest)
    return dest


def main() -> int:
    prose = (ROOT / "_prose.md").read_text()
    frontmatter = (ROOT / "_frontmatter.md").read_text()
    title_yml = (ROOT / "_title.yml").read_text()

    title = re.search(r'^title:\s*"(.+)"', title_yml, re.M).group(1)
    keywords = re.findall(r'^\s*-\s*"(.+)"', title_yml, re.M)

    sections = parse_sections(prose)
    by_head = {h: p for _, h, p in sections}

    def wc(head_names):
        n = 0
        started = False
        for lvl, head, paras in sections:
            if head in head_names:
                started = True
            if head == "Required end statements":
                started = False
            if started:
                for para in paras:
                    t = clean(para)
                    n += len(t.split())
        return n

    counts = {
        # Count the abstract body, not the five structural labels
        # ("Background." etc.), which is what the 250-word limit governs.
        "abstract": sum(
            len(re.sub(r"\*\*(Background|Aim|Methods|Results|Conclusion)\.\*\*\s*",
                       "", clean(p)).split())
            for p in by_head.get("Abstract", [])),
        "main": wc({"Introduction"}),
        "refs": len(re.findall(r"^\d+\. ", (ROOT / "_references.md").read_text(), re.M)),
        "keywords": keywords,
        "contributions": clean(" ".join(by_head.get("Authors' contributions", []))),
        "acknowledgements": clean(" ".join(by_head.get("Acknowledgements", []))),
        "data_availability": clean(" ".join(by_head.get("Data availability", []))),
    }

    OUT.mkdir(parents=True, exist_ok=True)
    files = [
        build_title_page(title, frontmatter, counts),
        build_manuscript(sections, title),
        build_ph_message(sections),
    ]
    for f in files:
        print(f"  {f.relative_to(ROOT)}  ({f.stat().st_size/1024:,.0f} KB)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
